import re
import pandas as pd
from typing import List, Dict, Tuple
import streamlit as st
import docx
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from PIL import Image

class ReferenceProcessor:
    def __init__(self):
        self.references = []
        self.reference_map = {}
        self.next_ref_id = 1
        
    def extract_references_from_text(self, text: str) -> Tuple[str, List[str]]:
        """Extrae referencias del texto y las reemplaza con marcadores numerados"""
        if not text:
            return text, []
            
        # Patrón para referencias en formato [[texto]]
        reference_pattern = r'\[\[([^\]]+)\]\]'
        references_found = re.findall(reference_pattern, text)
        
        if not references_found:
            return text, []
        
        processed_text = text
        replacement_refs = []
        
        for ref_content in references_found:
            # Separar referencias múltiples usando &&
            if '&&' in ref_content:
                individual_refs = [r.strip() for r in ref_content.split('&&')]
                ref_ids = []
                
                for individual_ref in individual_refs:
                    if individual_ref not in self.reference_map:
                        self.reference_map[individual_ref] = self.next_ref_id
                        self.references.append(individual_ref)
                        self.next_ref_id += 1
                    ref_ids.append(str(self.reference_map[individual_ref]))
                
                # Crear representación con comas en lugar de rangos
                compressed_ids = self._compress_number_ranges(ref_ids)
                replacement_text = f"[{compressed_ids}]"
                replacement_refs.extend(ref_ids)
                
            else:
                # Referencia simple
                if ref_content not in self.reference_map:
                    self.reference_map[ref_content] = self.next_ref_id
                    self.references.append(ref_content)
                    self.next_ref_id += 1
                
                ref_id = self.reference_map[ref_content]
                replacement_text = f"[{ref_id}]"
                replacement_refs.append(str(ref_id))
            
            processed_text = processed_text.replace(f"[[{ref_content}]]", replacement_text, 1)
        
        return processed_text, replacement_refs
    
    def _compress_number_ranges(self, number_list: List[str]) -> str:
        """Convierte una lista de números en lista con comas (sin rangos)"""
        if not number_list:
            return ""
        
        # Convertir a enteros y ordenar
        numbers = sorted(map(int, number_list))
        
        # Simplemente unir con comas, sin crear rangos
        return ", ".join(map(str, numbers))
    
    def extract_numeric_references_from_text(self, text: str) -> Tuple[str, List[str]]:
        """Extrae referencias numéricas existentes como [1], [2], etc. y las normaliza"""
        if not text:
            return text, []
            
        # Patrón para referencias numéricas como [1], [2], etc.
        numeric_ref_pattern = r'\[(\d+)\]'
        references_found = re.findall(numeric_ref_pattern, text)
        
        if not references_found:
            return text, []
        
        processed_text = text
        replacement_refs = []
        
        for ref_num in references_found:
            ref_id = int(ref_num)
            # Agregar a la lista de referencias si no existe
            ref_content = f"Referencia {ref_num}"
            if ref_content not in self.reference_map:
                self.reference_map[ref_content] = ref_id
                # Asegurarse de que la referencia esté en la lista
                while len(self.references) < ref_id:
                    self.references.append(f"Referencia {len(self.references) + 1}")
                if ref_id > len(self.references):
                    self.references.append(ref_content)
                else:
                    self.references[ref_id - 1] = ref_content
            
            replacement_refs.append(ref_num)
        
        return processed_text, replacement_refs
    
    def process_compressed_ranges(self, text: str) -> str:
        """Procesa y convierte rangos numéricos a listas con comas en el texto"""
        # Patrón para encontrar rangos como [1-3] o listas como [1,2,3]
        range_pattern = r'\[(\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*)\]'
        
        def expand_ranges(match):
            content = match.group(1)
            numbers = set()
            
            # Separar por comas
            parts = [part.strip() for part in content.split(',')]
            
            for part in parts:
                if '-' in part:
                    # Es un rango como 1-3
                    start, end = map(int, part.split('-'))
                    numbers.update(range(start, end + 1))
                else:
                    # Es un número individual
                    numbers.add(int(part))
            
            # Ordenar y convertir a lista con comas
            sorted_numbers = sorted(numbers)
            return f"[{', '.join(map(str, sorted_numbers))}]"
        
        return re.sub(range_pattern, expand_ranges, text)

class DOCXReferenceProcessor:
    def __init__(self):
        self.ref_processor = ReferenceProcessor()
    
    def process_docx_file(self, input_file, process_mode: str = "both") -> docx.Document:
        """Procesa un archivo DOCX manteniendo el formato original"""
        try:
            # Cargar documento
            doc = docx.Document(BytesIO(input_file.read()))
            
            # Procesar todo el documento en orden secuencial
            self._process_document_elements_sequential(doc, process_mode)
            
            return doc
            
        except Exception as e:
            st.error(f"Error al procesar el archivo DOCX: {e}")
            import traceback
            st.error(f"Detalle del error: {traceback.format_exc()}")
            return None
    
    def _process_document_elements_sequential(self, doc, process_mode: str):
        """Procesa todos los elementos del documento DOCX en orden secuencial"""
        
        # Procesar elementos en el orden exacto en que aparecen en el documento
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Es un párrafo
                paragraph = docx.text.paragraph.Paragraph(element, doc)
                if paragraph.text.strip():
                    processed_text = paragraph.text
                    
                    if process_mode in ["brackets", "both"]:
                        processed_text, _ = self.ref_processor.extract_references_from_text(processed_text)
                    
                    if process_mode in ["numeric", "both"]:
                        processed_text, _ = self.ref_processor.extract_numeric_references_from_text(processed_text)
                    
                    # Aplicar conversión de rangos numéricos a listas con comas
                    processed_text = self.ref_processor.process_compressed_ranges(processed_text)
                    
                    if paragraph.text != processed_text:
                        self._replace_paragraph_text_safe(paragraph, processed_text)
            
            elif element.tag.endswith('tbl'):  # Es una tabla
                table = docx.table.Table(element, doc)
                self._process_table_sequential(table, process_mode)
    
    def _process_table_sequential(self, table, process_mode: str):
        """Procesa una tabla en el orden de aparición"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        processed_text = paragraph.text
                        
                        if process_mode in ["brackets", "both"]:
                            processed_text, _ = self.ref_processor.extract_references_from_text(processed_text)
                        
                        if process_mode in ["numeric", "both"]:
                            processed_text, _ = self.ref_processor.extract_numeric_references_from_text(processed_text)
                        
                        # Aplicar conversión de rangos numéricos a listas con comas
                        processed_text = self.ref_processor.process_compressed_ranges(processed_text)
                        
                        if paragraph.text != processed_text:
                            self._replace_paragraph_text_safe(paragraph, processed_text)
    
    def _replace_paragraph_text_safe(self, paragraph, new_text):
        """Reemplaza el texto de un párrafo de manera segura manteniendo el formato"""
        if paragraph.text == new_text:
            return
            
        try:
            # Método seguro: limpiar el párrafo preservando el formato básico
            p = paragraph._p
            # Encontrar todos los elementos 'r' (runs) primero
            runs_to_remove = []
            for elem in p.iter():
                if elem.tag.endswith('r'):
                    runs_to_remove.append(elem)
            
            # Eliminar los runs de manera segura
            for run_elem in runs_to_remove:
                try:
                    p.remove(run_elem)
                except ValueError:
                    # Si el elemento ya no es hijo, continuar
                    continue
            
            # Agregar nuevo texto
            paragraph.add_run(new_text)
            
        except Exception as e:
            # Si falla el método anterior, usar método alternativo
            self._replace_paragraph_text_alternative(paragraph, new_text)
    
    def _replace_paragraph_text_alternative(self, paragraph, new_text):
        """Método alternativo para reemplazar texto usando clear()"""
        try:
            # Limpiar el párrafo usando clear()
            paragraph.clear()
            # Agregar nuevo texto
            paragraph.add_run(new_text)
        except Exception as e:
            # Último recurso: crear un nuevo párrafo
            parent = paragraph._p.getparent()
            if parent is not None:
                new_p = docx.oxml.OxmlElement('w:p')
                new_r = docx.oxml.OxmlElement('w:r')
                new_t = docx.oxml.OxmlElement('w:t')
                new_t.text = new_text
                new_r.append(new_t)
                new_p.append(new_r)
                parent.insert(parent.index(paragraph._p) + 1, new_p)
                parent.remove(paragraph._p)
    
    def generate_references_section(self, doc):
        """Genera la sección de referencias al final del documento"""
        if not self.ref_processor.references:
            return
        
        # Agregar salto de página
        doc.add_page_break()
        
        # Título de referencias
        title_paragraph = doc.add_paragraph("REFERENCES")
        title_paragraph.style = doc.styles['Heading1']
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar cada referencia
        for i, ref in enumerate(self.ref_processor.references, 1):
            ref_paragraph = doc.add_paragraph()
            ref_paragraph.style = doc.styles['Normal']
            ref_paragraph.add_run(f"{i}. {ref}")
    
    def get_statistics(self):
        """Obtiene estadísticas del procesamiento"""
        return {
            'total_references': len(self.ref_processor.references),
            'references_found': self.ref_processor.references,
            'reference_map': self.ref_processor.reference_map
        }

# Función para cargar y redimensionar el logo
def load_and_resize_logo(scale_factor=0.10):
    """Carga y redimensiona el logo del Instituto Nacional de Cardiología"""
    try:
        # Cargar la imagen del logo
        logo = Image.open("escudo_COLOR.jpg")
        original_width, original_height = logo.size
        
        # Calcular nuevas dimensiones con factor fijo de 0.10
        new_width = int(original_width * scale_factor)
        new_height = int(original_height * scale_factor)
        
        # Redimensionar el logo
        resized_logo = logo.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        return resized_logo, original_width, original_height
        
    except FileNotFoundError:
        st.warning("⚠️ No se encontró el archivo 'escudo_COLOR.jpg'. Asegúrate de que esté en el mismo directorio que el script.")
        return None, None, None
    except Exception as e:
        st.error(f"Error al cargar el logo: {e}")
        return None, None, None

# INTERFAZ STREAMLIT MEJORADA
def main():
    st.set_page_config(
        page_title="Procesador de Referencias DOCX - INC", 
        page_icon="📚", 
        layout="wide"
    )
    
    # Cargar y redimensionar el logo con factor fijo de 0.10
    logo, original_width, original_height = load_and_resize_logo(scale_factor=0.10)
    
    # Contenido principal - SIN columna izquierda
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Mostrar logo centrado en la parte superior
        if logo:
            st.image(logo, use_container_width=False)  # CORREGIDO: use_column_width → use_container_width
            # Mostrar el lema debajo del logo
            #st.markdown("""
            #<div style='text-align: center; font-style: italic; color: #666; margin-top: -10px; margin-bottom: 20px;'>
            #AMOR:SCIENTIA:OVE:INSERVIANT:CORDI<br>
            #INSTITUTO NACIONAL DE CARDIOLOGÍA IGNACIO CHÁVEZ
            #</div>
            #""", unsafe_allow_html=True)
        
        st.title("📚 Procesador de Referencias para Documentos DOCX")
        st.markdown("""
        Este programa procesa automáticamente las referencias bibliográficas en documentos Word, 
        manteniendo el formato original del documento.
        
        **Funcionalidades:**
        - ✅ **Referencia simple**:  `[[ref1]]` → `[1]`
        - ✅ **Referencias múltiples**: `[[ref1 && ref2 && ref3]]` → `[1, 2, 3]` (sin rangos)
        - ✅ **Procesamiento secuencial**: Las referencias se numeran en el orden exacto de aparición
        """)
    
    # SELECTOR DE MODO DE PROCESAMIENTO - SOLO UN MODO DISPONIBLE
    st.subheader("🔧 Configuración de Procesamiento")
    
    # Solo mostramos el modo de referencias en formato [[texto]]
    process_mode = "Referencias en formato [[texto]] → [número]"
    
    st.info(f"**Modo seleccionado:** {process_mode}")
    
    # Mapear selección a modo de procesamiento
    mode_map = {
        "Referencias en formato [[texto]] → [número]": "brackets"
    }
    
    # EJEMPLOS
    with st.expander("📖 Ver ejemplos de formato"):
        st.markdown("""
        **Delimitadores soportados:**
        - `[[Ref1]]` para encerrar las referencias
        - `&&` para separar múltiples referencias
        
        **Ejemplos:**
        - `[[Referencia simple]]` → `[1]`
        - `[[Ref1 && Ref2 && Ref3]]` → `[1, 2, 3]`
        """)
    
    uploaded_file = st.file_uploader("Sube tu archivo DOCX", type=['docx'])
    
    if uploaded_file is not None:
        # Mostrar información del archivo
        file_details = {
            "Nombre": uploaded_file.name,
            "Tipo": uploaded_file.type,
            "Tamaño": f"{uploaded_file.size / 1024:.2f} KB",
            "Modo de procesamiento": process_mode
        }
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Información del archivo")
            for key, value in file_details.items():
                st.write(f"**{key}:** {value}")
        
        # BOTÓN PARA PROCESAR
        st.markdown("---")
        col_process = st.columns([1, 2, 1])
        with col_process[1]:
            process_btn = st.button(
                "🚀 **PROCESAR REFERENCIAS**", 
                type="primary", 
                use_container_width=True,
                help="Haz clic para procesar las referencias del documento"
            )
        
        # Inicializar variables de sesión si no existen
        if 'processed_doc' not in st.session_state:
            st.session_state.processed_doc = None
        if 'stats' not in st.session_state:
            st.session_state.stats = None
        if 'output_buffer' not in st.session_state:
            st.session_state.output_buffer = None
        if 'text_content' not in st.session_state:
            st.session_state.text_content = None
        
        if process_btn:
            # Procesar documento
            processor = DOCXReferenceProcessor()
            
            with st.spinner("Procesando documento DOCX..."):
                processed_doc = processor.process_docx_file(uploaded_file, mode_map[process_mode])
            
            if processed_doc:
                # Generar sección de referencias
                processor.generate_references_section(processed_doc)
                
                # Obtener estadísticas
                stats = processor.get_statistics()
                
                # Preparar archivo para descarga
                output_buffer = BytesIO()
                processed_doc.save(output_buffer)
                output_buffer.seek(0)
                
                # Preparar contenido de texto para referencia
                text_content = "Documento procesado con referencias numeradas\n\n"
                text_content += f"Total de referencias: {stats['total_references']}\n\n"
                
                for i, ref in enumerate(stats['references_found'], 1):
                    text_content += f"{i}. {ref}\n"
                
                # Guardar en session state
                st.session_state.processed_doc = processed_doc
                st.session_state.stats = stats
                st.session_state.output_buffer = output_buffer
                st.session_state.text_content = text_content
        
        # Mostrar resultados si existen en session state
        if st.session_state.stats is not None:
            stats = st.session_state.stats
            
            with col2:
                st.subheader("📊 Estadísticas de procesamiento")
                st.metric("Referencias encontradas", stats['total_references'])
                
                if stats['total_references'] > 0:
                    st.success("✅ Referencias procesadas exitosamente!")
                    # Mostrar preview de referencias
                    st.subheader("🔍 Referencias detectadas")
                    
                    with st.expander("Ver lista completa de referencias"):
                        for i, ref in enumerate(stats['references_found'], 1):
                            st.write(f"**[{i}]** {ref}")
                else:
                    st.warning("⚠️ No se encontraron referencias en el documento")
            
            # Botones de descarga (siempre visibles una vez procesado)
            st.subheader("💾 Descargar documento procesado")
            
            col_d1, col_d2 = st.columns(2)
            
            with col_d1:
                st.download_button(
                    label="📥 Descargar DOCX procesado",
                    data=st.session_state.output_buffer.getvalue() if st.session_state.output_buffer else "",
                    file_name=f"procesado_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx"
                )
            
            with col_d2:
                st.download_button(
                    label="📄 Descargar lista de referencias (TXT)",
                    data=st.session_state.text_content if st.session_state.text_content else "",
                    file_name="referencias_procesadas.txt",
                    mime="text/plain",
                    key="download_txt"
                )
            
            # Mostrar información adicional
            st.info("""
            **Nota:** El documento descargado mantiene todo el formato original de Word, 
            con las referencias convertidas a numeración consistente y una sección 
            REFERENCES agregada al final.
            
            **Mejora importante:** 
            - Las referencias se numeran en el orden exacto de aparición en el documento
            - Todas las referencias múltiples se muestran como listas con comas: `[1, 2, 3]`
            """)
        
        elif not process_btn:
            st.info("👆 **Haz clic en el botón 'PROCESAR REFERENCIAS' para comenzar**")

if __name__ == "__main__":
    main()
